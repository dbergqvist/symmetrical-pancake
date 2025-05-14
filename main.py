#!/usr/bin/env python3

import os
import random
import string
import json
import time
import numpy as np
from concurrent.futures import ProcessPoolExecutor
from tqdm import tqdm

# Document generation libraries
import docx
import pandas as pd
from fpdf import FPDF
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference

# For text generation (if not using an external API)
import nltk
from nltk.corpus import gutenberg

# Config
OUTPUT_DIR = "synthetic_docs"
NUM_DOCUMENTS = 10  # One million documents
DOCUMENT_TYPES = {
    "docx": 0.4,  # 40% Word documents
    "pdf": 0.3,   # 30% PDFs
    "xlsx": 0.2,  # 20% Excel spreadsheets
    "txt": 0.1    # 10% Text files
}

# Content templates and data sources
DOCUMENT_TEMPLATES = {
    "report": 0.3,
    "letter": 0.2,
    "memo": 0.15,
    "invoice": 0.15,
    "data_analysis": 0.2
}

# Setup
os.makedirs(OUTPUT_DIR, exist_ok=True)
nltk.download('gutenberg')

# Text generation functions
def generate_paragraphs(num_paragraphs=3, sentences_per_paragraph=5):
    """Generate random paragraphs of text"""
    try:
        # Ensure NLTK data is downloaded
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            print("Downloading NLTK punkt tokenizer...")
            nltk.download('punkt', quiet=True)
        
        try:
            nltk.data.find('corpora/gutenberg')
        except LookupError:
            print("Downloading NLTK Gutenberg corpus...")
            nltk.download('gutenberg', quiet=True)
        
        # Get raw text from multiple books
        books = ['austen-sense.txt', 'austen-emma.txt', 'austen-persuasion.txt', 
                'bible-kjv.txt', 'blake-poems.txt', 'carroll-alice.txt']
        
        all_text = []
        for book in books:
            try:
                text = gutenberg.raw(book)
                # Split into sentences using NLTK's sentence tokenizer
                sentences = nltk.sent_tokenize(text)
                all_text.extend(sentences)
            except Exception as e:
                print(f"Warning: Could not load {book}: {str(e)}")
                continue
        
        if not all_text:
            raise Exception("No text could be loaded from any book")
        
        paragraphs = []
        for _ in range(num_paragraphs):
            selected_sentences = random.sample(all_text, min(sentences_per_paragraph, len(all_text)))
            paragraph = ' '.join(selected_sentences)
            paragraphs.append(paragraph)
        
        return paragraphs
    except Exception as e:
        print(f"Warning: Using fallback text due to error: {str(e)}")
        # Fallback text if NLTK data is not available
        fallback_text = [
            "This is a sample paragraph for testing purposes. It contains multiple sentences that form a coherent text block.",
            "The quick brown fox jumps over the lazy dog. This is a common pangram used for testing text rendering.",
            "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.",
            "In a world of digital transformation, data analysis plays a crucial role in decision-making processes.",
            "The importance of clear communication cannot be overstated in professional settings.",
            "Data-driven insights have become essential for modern business operations and strategic planning.",
            "Effective communication is the cornerstone of successful project management and team collaboration.",
            "Innovation and adaptability are key factors in maintaining competitive advantage in today's market.",
            "Quality assurance processes ensure consistent delivery of products and services to customers.",
            "Strategic planning and execution are vital for achieving long-term organizational goals.",
            "The integration of artificial intelligence and machine learning is revolutionizing business processes.",
            "Sustainable practices and environmental responsibility are increasingly important in corporate strategy.",
            "Customer experience and satisfaction remain top priorities in service-oriented industries.",
            "Digital transformation initiatives require careful planning and stakeholder engagement.",
            "Risk management and compliance frameworks are essential for organizational resilience."
        ]
        return random.sample(fallback_text, min(num_paragraphs, len(fallback_text)))

# Document generation functions
def create_word_document(filename, template_type):
    """Create a Word document"""
    doc = docx.Document()
    
    # Add a title
    doc.add_heading(f"{template_type.title()} - {' '.join(random.sample(string.ascii_uppercase, 3))}", 0)
    
    # Add content based on template type
    if template_type == "report":
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(generate_paragraphs(1)[0])
        
        doc.add_heading("Introduction", level=1)
        for paragraph in generate_paragraphs(2):
            doc.add_paragraph(paragraph)
            
        doc.add_heading("Findings", level=1)
        for paragraph in generate_paragraphs(3):
            doc.add_paragraph(paragraph)
            
        doc.add_heading("Conclusion", level=1)
        doc.add_paragraph(generate_paragraphs(1)[0])
        
    elif template_type == "letter":
        # Date
        doc.add_paragraph(f"May {random.randint(1, 31)}, 2025")
        doc.add_paragraph("")
        
        # Recipient
        doc.add_paragraph(f"Dear Mr./Ms. {random.choice(string.ascii_uppercase)}.,")
        
        # Body
        for paragraph in generate_paragraphs(3):
            doc.add_paragraph(paragraph)
            
        # Sign-off
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph(f"{random.choice(['John', 'Jane', 'Alex', 'Sam'])} {random.choice(string.ascii_uppercase)}.")
        
    # Save the document
    doc.save(filename)

def create_excel_document(filename, template_type):
    """Create an Excel spreadsheet"""
    wb = Workbook()
    ws = wb.active
    
    if template_type == "data_analysis":
        # Header row
        headers = ["Region", "Q1", "Q2", "Q3", "Q4", "Total"]
        for col, header in enumerate(headers, 1):
            ws.cell(row=1, column=col, value=header)
        
        # Data rows
        regions = ["North", "South", "East", "West", "Central"]
        row = 2
        
        for region in regions:
            quarterly_data = [random.randint(100, 1000) for _ in range(4)]
            total = sum(quarterly_data)
            
            ws.cell(row=row, column=1, value=region)
            for col, value in enumerate(quarterly_data, 2):
                ws.cell(row=row, column=col, value=value)
            ws.cell(row=row, column=6, value=total)
            
            row += 1
        
        # Add a chart
        chart = BarChart()
        chart.title = "Quarterly Performance by Region"
        chart.x_axis.title = "Region"
        chart.y_axis.title = "Revenue"
        
        data = Reference(ws, min_col=2, min_row=1, max_row=6, max_col=5)
        cats = Reference(ws, min_col=1, min_row=2, max_row=6)
        chart.add_data(data, titles_from_data=True)
        chart.set_categories(cats)
        
        ws.add_chart(chart, "H2")
        
    elif template_type == "invoice":
        # Header
        ws["A1"] = "INVOICE"
        ws["A3"] = "Bill To:"
        ws["A4"] = f"Company {random.choice(string.ascii_uppercase)}"
        ws["A5"] = f"{random.randint(100, 999)} Main Street"
        ws["A6"] = f"City, State {random.randint(10000, 99999)}"
        
        # Invoice details
        ws["E3"] = "Invoice #:"
        ws["F3"] = f"INV-{random.randint(1000, 9999)}"
        ws["E4"] = "Date:"
        ws["F4"] = f"5/{random.randint(1, 31)}/2025"
        ws["E5"] = "Due Date:"
        ws["F5"] = f"6/{random.randint(1, 30)}/2025"
        
        # Items table
        headers = ["Item", "Description", "Quantity", "Unit Price", "Amount"]
        for col, header in enumerate(headers, 1):
            ws.cell(row=8, column=col, value=header)
        
        # Items
        num_items = random.randint(3, 8)
        row = 9
        total = 0
        
        for i in range(num_items):
            item = f"Item {random.choice(string.ascii_uppercase)}-{random.randint(100, 999)}"
            desc = f"Description for {item}"
            qty = random.randint(1, 10)
            price = round(random.uniform(10, 500), 2)
            amount = qty * price
            total += amount
            
            ws.cell(row=row, column=1, value=item)
            ws.cell(row=row, column=2, value=desc)
            ws.cell(row=row, column=3, value=qty)
            ws.cell(row=row, column=4, value=price)
            ws.cell(row=row, column=5, value=amount)
            
            row += 1
        
        # Total
        ws.cell(row=row+1, column=4, value="Total:")
        ws.cell(row=row+1, column=5, value=total)
    
    wb.save(filename)

def create_pdf_document(filename, template_type):
    """Create a PDF document"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    if template_type == "memo":
        # Header
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, "MEMORANDUM", ln=True, align='C')
        pdf.ln(10)
        
        # Memo details
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(30, 10, "TO:", ln=0)
        pdf.set_font("Arial", '', 12)
        pdf.cell(160, 10, f"All {random.choice(['Employees', 'Managers', 'Department Heads', 'Team Members'])}", ln=1)
        
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(30, 10, "FROM:", ln=0)
        pdf.set_font("Arial", '', 12)
        pdf.cell(160, 10, f"{random.choice(['John', 'Jane', 'Alex', 'Sam'])} {random.choice(string.ascii_uppercase)}.", ln=1)
        
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(30, 10, "DATE:", ln=0)
        pdf.set_font("Arial", '', 12)
        pdf.cell(160, 10, f"May {random.randint(1, 31)}, 2025", ln=1)
        
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(30, 10, "SUBJECT:", ln=0)
        pdf.set_font("Arial", '', 12)
        pdf.cell(160, 10, f"{random.choice(['Policy Update', 'Upcoming Event', 'Quarterly Results', 'New Initiative'])}", ln=1)
        
        pdf.ln(10)
        
        # Content
        for paragraph in generate_paragraphs(4):
            pdf.multi_cell(0, 10, paragraph)
            pdf.ln(5)
    
    elif template_type == "report":
        # Title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, f"REPORT: {' '.join(random.sample(string.ascii_uppercase, 3))}", ln=True, align='C')
        pdf.ln(10)
        
        sections = ["Executive Summary", "Introduction", "Methodology", "Findings", "Conclusion"]
        
        for section in sections:
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(200, 10, section, ln=True)
            pdf.set_font("Arial", '', 12)
            
            for paragraph in generate_paragraphs(random.randint(1, 3)):
                pdf.multi_cell(0, 10, paragraph)
                pdf.ln(5)
    
    pdf.output(filename)

def create_text_document(filename, template_type):
    """Create a simple text file"""
    with open(filename, 'w') as f:
        if template_type == "memo":
            f.write("MEMORANDUM\n\n")
            f.write(f"TO: All {random.choice(['Employees', 'Managers', 'Department Heads', 'Team Members'])}\n")
            f.write(f"FROM: {random.choice(['John', 'Jane', 'Alex', 'Sam'])} {random.choice(string.ascii_uppercase)}.\n")
            f.write(f"DATE: May {random.randint(1, 31)}, 2025\n")
            f.write(f"SUBJECT: {random.choice(['Policy Update', 'Upcoming Event', 'Quarterly Results', 'New Initiative'])}\n\n")
            
            for paragraph in generate_paragraphs(4):
                f.write(paragraph + "\n\n")
        
        else:  # generic text file
            f.write(f"DOCUMENT: {' '.join(random.sample(string.ascii_uppercase, 3))}\n\n")
            for paragraph in generate_paragraphs(random.randint(5, 10)):
                f.write(paragraph + "\n\n")

def generate_document(doc_index):
    """Generate a single document based on probabilities"""
    # Determine document type
    doc_type = random.choices(
        list(DOCUMENT_TYPES.keys()),
        weights=list(DOCUMENT_TYPES.values()),
        k=1
    )[0]
    
    # Determine template type
    template_type = random.choices(
        list(DOCUMENT_TEMPLATES.keys()),
        weights=list(DOCUMENT_TEMPLATES.values()),
        k=1
    )[0]
    
    # Create filename
    rand_id = ''.join(random.choices(string.ascii_lowercase + string.digits, k=8))
    filename = os.path.join(OUTPUT_DIR, f"doc_{doc_index:07d}_{template_type}_{rand_id}.{doc_type}")
    
    # Create document based on type
    try:
        if doc_type == "docx":
            create_word_document(filename, template_type)
        elif doc_type == "xlsx":
            create_excel_document(filename, template_type)
        elif doc_type == "pdf":
            create_pdf_document(filename, template_type)
        elif doc_type == "txt":
            create_text_document(filename, template_type)
            
        return {"index": doc_index, "filename": filename, "type": doc_type, "template": template_type, "status": "success"}
    except Exception as e:
        return {"index": doc_index, "error": str(e), "status": "failed"}

def main():
    """Main function to generate documents in parallel"""
    print(f"Generating {NUM_DOCUMENTS} synthetic documents...")
    start_time = time.time()
    
    # Generate document indices
    doc_indices = list(range(1, NUM_DOCUMENTS + 1))
    
    # Process in batches to avoid memory issues
    batch_size = 10000
    total_batches = (NUM_DOCUMENTS + batch_size - 1) // batch_size
    
    results = []
    
    for batch in tqdm(range(total_batches), desc="Processing batches"):
        batch_start = batch * batch_size
        batch_end = min((batch + 1) * batch_size, NUM_DOCUMENTS)
        batch_indices = doc_indices[batch_start:batch_end]
        
        # Process batch in parallel
        with ProcessPoolExecutor(max_workers=os.cpu_count()) as executor:
            batch_results = list(tqdm(
                executor.map(generate_document, batch_indices),
                total=len(batch_indices),
                desc=f"Batch {batch+1}/{total_batches}"
            ))
            
            results.extend(batch_results)
        
        # Optional: Save progress periodically
        if (batch + 1) % 10 == 0 or batch == total_batches - 1:
            with open(os.path.join(OUTPUT_DIR, f"generation_report_{batch+1}.json"), 'w') as f:
                json.dump({
                    "completed_batches": batch + 1,
                    "total_batches": total_batches,
                    "documents_generated": len(results),
                    "success_count": sum(1 for r in results if r["status"] == "success"),
                    "failed_count": sum(1 for r in results if r["status"] == "failed"),
                }, f, indent=2)
    
    # Calculate statistics
    end_time = time.time()
    total_time = end_time - start_time
    success_count = sum(1 for r in results if r["status"] == "success")
    
    # Save final report
    with open(os.path.join(OUTPUT_DIR, "generation_report_final.json"), 'w') as f:
        json.dump({
            "total_documents": NUM_DOCUMENTS,
            "success_count": success_count,
            "failed_count": NUM_DOCUMENTS - success_count,
            "time_taken_seconds": total_time,
            "documents_per_second": NUM_DOCUMENTS / total_time,
            "document_type_distribution": {
                doc_type: sum(1 for r in results if r.get("status") == "success" and r.get("type") == doc_type) / success_count
                for doc_type in DOCUMENT_TYPES
            },
            "template_type_distribution": {
                template_type: sum(1 for r in results if r.get("status") == "success" and r.get("template") == template_type) / success_count
                for template_type in DOCUMENT_TEMPLATES
            }
        }, f, indent=2)
    
    print(f"Document generation complete!")
    print(f"Generated {success_count} documents in {total_time:.2f} seconds")
    print(f"Average: {success_count / total_time:.2f} documents per second")
    print(f"Final report saved to {os.path.join(OUTPUT_DIR, 'generation_report_final.json')}")

if __name__ == "__main__":
    main()